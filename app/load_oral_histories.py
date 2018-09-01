import os, re
from glob import glob
import csv

from app import basedir

from docx import Document

ORAL_HISTORY_FNAMES = glob(os.path.join(basedir,  'static/LSF_oral_histories/*'))
ORAL_HISTORY_FNAMES_INTERVIEWEE_MATCH_FILE = 'app/static/LSF_oral_histories/oral_histories_fname_interviewee_match.tsv'

# Captures, e.g. ('(INTERVIEWEE:) (text...)') or ('(DATE:) (<date>)')
PATTERN_LINEBEGIN = re.compile(r"^([A-Z]\S*?:)\s(.*)")

PATTERN_BRACKETS_NOSPACEBEFORE = re.compile(r"(?<=\S)(\[.*?\])")
PATTERN_BRACKETS_NOSPACEAFTER = re.compile(r"(\[.*?\])(?=\S)")

def load_oral_histories_fnames():
    oral_histories_basenames = [(fname, os.path.basename(fname)) for fname in ORAL_HISTORY_FNAMES]
    with open(ORAL_HISTORY_FNAMES_INTERVIEWEE_MATCH_FILE, 'r') as f:
        reader = csv.DictReader(f, delimiter='\t')
        # match to filename in static folder:
        for r in reader:
            for fname, fname_base in oral_histories_basenames:
                if r['fname_base'] == fname_base:
                    break
            yield {
                'fname': fname,
                'fname_base': r['fname_base'],
                'interviewee_full_name': r['interviewee_full_name'],
                'interviewee_last_name': r['interviewee_last_name'],
                'interviewee_first_name': r['interviewee_first_name'],
                }

def preprocess_oral_history(text):
    """Take in the document as a series of paragraphs (lines), apply some minimal processing, and output a list of lines of text

    :text: list of strings
    :returns: list of strings

    """
    flag = False
    lines = []
    for p in text:
        line = p.strip()
        if line:
            # try to find a certain kind of line that tends to mark the beginning of the oral history text:
            # (this helps with tokenization later)
            m = PATTERN_LINEBEGIN.search(line)
            if m:
                if m.group(1).lower().startswith('date'):
                    flag = True
                    continue
            if flag is True:
                # add spaces before and after brackets, if there aren't any
                line = re.sub(PATTERN_BRACKETS_NOSPACEBEFORE, r" \1", line)
                line = re.sub(PATTERN_BRACKETS_NOSPACEAFTER, r"\1 ", line)
                lines.append(line)
    if len(lines) == 0:
        # the processing failed. fall back to this:
        lines = [p for p in text]
    return lines

# def preprocess_oral_history(document):
#     """TODO: Docstring for preprocess_oral_history.
#
#     :document: python-docx document
#     :returns: TODO
#
#     """
#     flag = False
#     lines = []
#     for p in document.paragraphs:
#         line = p.text.strip()
#         if line:
#             m = PATTERN_LINEBEGIN.search(line)
#             if m:
#                 if m.group(1).lower().startswith('date'):
#                     flag = True
#                     continue
#             if flag is True:
#                 lines.append(line)
#     if len(lines) == 0:
#         # the processing failed. fall back to this:
#         lines = [p.text for p in document.paragraphs]
#
#     return lines

def test():
    test_history = 'static/LSF_oral_histories/Addario AE-GP.docx'
    test_history = os.path.join(basedir, test_history)
    # data = os.path.exists(test_history)
    document = Document(test_history)
    data = [p.text for p in document.paragraphs]
    return data
